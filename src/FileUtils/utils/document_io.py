from __future__ import annotations

from pathlib import Path
from typing import Dict, Any, Union


def save_markdown(
    content: Union[str, Dict[str, Any]], path: Path, *, encoding: str = "utf-8"
) -> str:
    if isinstance(content, str):
        markdown_content = content
    elif isinstance(content, dict):
        frontmatter = content.get("frontmatter", {})
        body = content.get("body", "")
        if frontmatter:
            import yaml

            frontmatter_yaml = yaml.safe_dump(frontmatter, default_flow_style=False)
            markdown_content = f"---\n{frontmatter_yaml}---\n\n{body}"
        else:
            markdown_content = body
    else:
        markdown_content = str(content)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding=encoding) as f:
        f.write(markdown_content)
    return str(path)


def load_markdown(path: Path, *, encoding: str = "utf-8") -> Union[str, Dict[str, Any]]:
    with open(path, "r", encoding=encoding) as f:
        content = f.read()
    if content.startswith("---\n"):
        try:
            import yaml

            parts = content.split("---\n", 2)
            if len(parts) >= 3:
                frontmatter = yaml.safe_load(parts[1])
                body = parts[2].strip()
                return {"frontmatter": frontmatter or {}, "body": body}
        except Exception:
            pass
    return content


def save_docx_simple(content: Union[str, Dict[str, Any]], path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if isinstance(content, str):
        doc.add_paragraph(content)
    elif isinstance(content, dict):
        if "title" in content:
            doc.add_heading(content["title"], 0)
        for section in content.get("sections", []):
            if "heading" in section:
                doc.add_heading(section["heading"], level=section.get("level", 1))
            if "text" in section:
                doc.add_paragraph(section["text"])
            if (
                "table" in section
                and isinstance(section["table"], list)
                and section["table"]
            ):
                table = doc.add_table(
                    rows=len(section["table"]), cols=len(section["table"][0])
                )
                for i, row in enumerate(section["table"]):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = str(cell)
    else:
        doc.add_paragraph(str(content))
    doc.save(path)
    return str(path)


def load_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e
    doc = Document(path)
    text: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text.append(" | ".join(cells))
    return "\n".join(text)


def save_pdf_text(content: Union[str, Dict[str, Any]], path: Path) -> str:
    try:
        import fitz
    except ImportError as e:
        raise RuntimeError("PyMuPDF not installed") from e
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    page = doc.new_page()
    if isinstance(content, str):
        page.insert_text((50, 50), content, fontsize=12)
    elif isinstance(content, dict):
        y = 50
        if "title" in content:
            page.insert_text((50, y), content["title"], fontsize=16)
            y += 30
        for section in content.get("sections", []):
            if "heading" in section:
                page.insert_text((50, y), section["heading"], fontsize=14)
                y += 25
            if "text" in section:
                page.insert_text((50, y), section["text"], fontsize=12)
                y += 20
    else:
        page.insert_text((50, 50), str(content), fontsize=12)
    doc.save(path)
    doc.close()
    return str(path)


def load_pdf_text(path: Path) -> str:
    try:
        import fitz
    except ImportError as e:
        raise RuntimeError("PyMuPDF not installed") from e
    doc = fitz.open(path)
    text = []
    for i in range(doc.page_count):
        pg = doc[i]
        t = pg.get_text()
        if t.strip():
            text.append(t)
    doc.close()
    return "\n\n".join(text)


def save_pptx(content: Union[bytes, Path, str], path: Path) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(content, bytes):
        path.write_bytes(content)
        return str(path)
    src = Path(content)
    if not src.exists():
        raise RuntimeError("Source PPTX file not found")
    path.write_bytes(src.read_bytes())
    return str(path)
