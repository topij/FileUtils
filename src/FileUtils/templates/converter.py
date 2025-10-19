"""Markdown to DOCX Converter for FileUtils.

Enhanced markdown conversion with template support and comprehensive formatting.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, Union, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..utils.common import get_logger
from .styles import StyleMapper


class MarkdownToDocxConverter:
    """Comprehensive markdown to DOCX converter with template support."""
    
    def __init__(self, template_manager, style_mapper: Optional[StyleMapper] = None):
        """Initialize converter.
        
        Args:
            template_manager: DocxTemplateManager instance
            style_mapper: StyleMapper instance
        """
        self.template_manager = template_manager
        self.style_mapper = style_mapper or StyleMapper()
        self.logger = get_logger(self.__class__.__name__)
    
    def convert_markdown_to_docx(
        self,
        markdown_content: str,
        template_name: Optional[str] = None,
        add_provenance: bool = True,
        add_reviewer_instructions: bool = False,
        source_file: Optional[str] = None,
        **kwargs
    ) -> Document:
        """Convert markdown content to DOCX document.
        
        Args:
            markdown_content: Markdown content to convert
            template_name: Name of template to use
            add_provenance: Whether to add provenance header
            add_reviewer_instructions: Whether to add reviewer instructions
            source_file: Source file path for provenance
            **kwargs: Additional options
            
        Returns:
            DOCX Document object
        """
        # Create document from template or default
        doc = self._create_document_from_template(template_name)
        
        # Add reviewer instructions if requested
        if add_reviewer_instructions:
            self._add_reviewer_instructions(doc)
        
        # Add provenance header if requested
        if add_provenance:
            self._add_provenance_header(doc, source_file)
        
        # Convert markdown content
        self._convert_content(doc, markdown_content)
        
        return doc
    
    def _create_document_from_template(self, template_name: Optional[str] = None) -> Document:
        """Create document from template or default."""
        template_path = self.template_manager.get_template_path(template_name)
        
        if template_path and template_path.exists():
            try:
                doc = Document(template_path)
                # Clear template content while preserving styles
                self._clear_template_content(doc)
                self.logger.debug(f"Created document from template: {template_path}")
                return doc
            except Exception as e:
                self.logger.warning(f"Failed to load template {template_path}: {e}")
        
        # Fallback to default document
        self.logger.debug("Using default document template")
        return Document()
    
    def _clear_template_content(self, doc: Document):
        """Clear template content while preserving styles."""
        # Remove all paragraphs
        while len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Remove all tables
        while len(doc.tables) > 0:
            t = doc.tables[0]._element
            t.getparent().remove(t)
        
        # Clear document body
        try:
            body = doc._body
            for element in list(body):
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    body.remove(element)
        except Exception:
            pass
    
    def _add_reviewer_instructions(self, doc: Document):
        """Add reviewer instructions section."""
        # Add main heading
        heading = doc.add_heading("Reviewer Instructions", level=1)
        
        # Add instructions paragraph
        instructions = doc.add_paragraph()
        instructions.add_run("This document contains TODO items and other content that needs your review. You can:")
        
        # Add numbered list of instructions
        instructions_list = [
            "Review and resolve TODO items marked with <!-- TODO: ... -->",
            "Add RESOLUTION field to each TODO: <!-- TODO: [description]. RESOLUTION: [done | not relevant | tbd] -->",
            "When fully resolved, change <!-- TODO: to <!-- RESOLVED: and add resolution notes",
            "Modify, add, or comment on any other content beyond TODO items",
            "Update document content based on your review decisions",
            "Return the document with your changes and resolutions marked",
            "NOTE: You can also mark/comment any of the items as irrelevant or not needed. Or suggest moving it to some other place/phase"
        ]
        
        for i, instruction in enumerate(instructions_list, 1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. ").bold = True
            para.add_run(instruction)
            # Apply list style
            self.style_mapper.apply_style_safely(para, self.style_mapper.get_style("list_number"))
        
        # Add separator
        doc.add_paragraph()
        separator = doc.add_paragraph("─" * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def _add_provenance_header(self, doc: Document, source_file: Optional[str] = None):
        """Add provenance header."""
        if source_file:
            filename = Path(source_file).name
        else:
            filename = "markdown content"
        
        current_date = datetime.now().strftime("%Y-%m-%d")
        provenance_text = f"Autogenerated from {filename} on {current_date}"
        
        # Add as centered paragraph
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Style the provenance text
        run = para.add_run(provenance_text)
        run.italic = True
        run.font.size = Pt(9)
        
        # Add empty line after provenance
        doc.add_paragraph()
    
    def _convert_content(self, doc: Document, content: str):
        """Convert markdown content to DOCX elements."""
        # Remove frontmatter
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                content = parts[2].strip()
        
        # Split into lines and clean up
        lines = content.split('\n')
        
        # Remove leading empty lines
        while lines and not lines[0].strip():
            lines.pop(0)
        
        # Remove lines that are just three dashes
        lines = [line for line in lines if line.strip() != '---']
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                self._add_heading(doc, text, level)
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                i = self._add_bullet_list(doc, lines, i)
                continue
            
            # Numbered lists
            elif re.match(r'^\d+\.\s', line):
                i = self._add_numbered_list(doc, lines, i)
                continue
            
            # Tables
            elif self._is_table_line(line) and i + 1 < len(lines):
                i = self._add_table(doc, lines, i)
                continue
            
            # Regular paragraphs
            else:
                self._add_paragraph(doc, line)
            
            i += 1
    
    def _is_table_line(self, line: str) -> bool:
        """Check if a line looks like a table row."""
        return '|' in line and len(line.split('|')) >= 3
    
    def _is_table_separator(self, line: str) -> bool:
        """Check if a line is a table separator."""
        line = line.strip()
        if not '|' in line:
            return False
        
        parts = line.split('|')
        if len(parts) < 3:
            return False
        
        separator_count = 0
        for part in parts:
            part = part.strip()
            if part and re.match(r'^[-:]+$', part):
                separator_count += 1
        
        return separator_count >= len([p for p in parts if p.strip()]) * 0.8
    
    def _add_heading(self, doc: Document, text: str, level: int):
        """Add a heading with formatted text."""
        heading = doc.add_heading(level=level)
        self._add_formatted_text(heading, text)
        return heading
    
    def _add_bullet_list(self, doc: Document, lines: List[str], start_index: int) -> int:
        """Add a bullet list with checkbox support."""
        list_items = []
        i = start_index
        
        # Collect all list items
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('- ') or line.startswith('* '):
                list_items.append(line[2:].strip())
                i += 1
            else:
                break
        
        # Add list items
        for item in list_items:
            para = doc.add_paragraph()
            
            # Check for checkbox syntax
            checkbox_match = re.match(r'^\s*\[\s*([xX]?)\s*\]\s*(.*)', item)
            if checkbox_match:
                is_checked = checkbox_match.group(1).lower() == 'x'
                checkbox_text = checkbox_match.group(2).strip()
                
                # Add checkbox
                run = para.add_run()
                run.add_text('☐' if not is_checked else '☑')
                run.font.name = 'Segoe UI Symbol'
                run.font.size = Pt(12)
                
                # Add space and process checkbox text
                para.add_run(' ')
                self._add_formatted_text_single(para, checkbox_text)
                
                # Apply list paragraph style
                self.style_mapper.apply_style_safely(para, self.style_mapper.get_style("list_paragraph"))
            else:
                # Regular bullet item
                self._add_formatted_text(para, item)
                self.style_mapper.apply_style_safely(para, self.style_mapper.get_style("list_bullet"))
        
        return i
    
    def _add_numbered_list(self, doc: Document, lines: List[str], start_index: int) -> int:
        """Add a numbered list."""
        list_items = []
        i = start_index
        
        # Collect all list items
        while i < len(lines):
            line = lines[i].strip()
            if re.match(r'^\d+\.\s', line):
                # Remove number prefix
                item = re.sub(r'^\d+\.\s', '', line)
                list_items.append(item)
                i += 1
            else:
                break
        
        # Add list items
        for item in list_items:
            para = doc.add_paragraph()
            self._add_formatted_text(para, item)
            self.style_mapper.apply_style_safely(para, self.style_mapper.get_style("list_number"))
        
        return i
    
    def _add_table(self, doc: Document, lines: List[str], start_index: int) -> int:
        """Add a table with proper styling."""
        table_lines = []
        i = start_index
        
        # Collect table lines, skip separators
        while i < len(lines):
            line = lines[i].strip()
            
            if self._is_table_separator(line):
                i += 1
                continue
            elif self._is_table_line(line):
                table_lines.append(line)
                i += 1
            else:
                break
        
        if len(table_lines) < 1:
            return start_index + 1
        
        # Parse table
        rows = []
        for line in table_lines:
            if '|' in line:
                # Split by | and clean up
                cells = [cell.strip() for cell in line.split('|')]
                # Remove empty cells at start/end
                if cells and not cells[0]:
                    cells = cells[1:]
                if cells and not cells[-1]:
                    cells = cells[:-1]
                rows.append(cells)
        
        if len(rows) < 1:
            return start_index + 1
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        
        # Apply table style with fallback
        table_style_chain = self.style_mapper.get_table_style_chain()
        for style_name in table_style_chain:
            if style_name:
                try:
                    table.style = style_name
                    break
                except Exception:
                    continue
        
        # Fill table data
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if i < len(table.rows) and j < len(table.rows[i].cells):
                    # Add formatted text to cell
                    cell_para = table.rows[i].cells[j].paragraphs[0]
                    self._add_formatted_text(cell_para, cell)
        
        return start_index + len(table_lines) + 1  # +1 for separator line
    
    def _add_paragraph(self, doc: Document, text: str):
        """Add a regular paragraph with formatting."""
        if text.strip():
            para = doc.add_paragraph()
            self._add_formatted_text(para, text)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to a paragraph."""
        # Process markdown links first: [text](link) -> text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Handle <br> tags
        if '<br>' in text or '<br/>' in text or '<br />' in text:
            text_parts = re.split(r'<br\s*/?>', text)
            for i, part in enumerate(text_parts):
                if part.strip():
                    if i > 0:
                        paragraph.add_run().add_break()
                    self._add_formatted_text_single(paragraph, part.strip())
            return
        
        # Process normally
        self._add_formatted_text_single(paragraph, text)
    
    def _add_formatted_text_single(self, paragraph, text: str):
        """Add formatted text without br tag processing."""
        # Check for checkbox syntax
        checkbox_match = re.match(r'^\s*\[\s*([xX]?)\s*\]\s*(.*)', text)
        if checkbox_match:
            is_checked = checkbox_match.group(1).lower() == 'x'
            checkbox_text = checkbox_match.group(2).strip()
            
            # Add checkbox
            run = paragraph.add_run()
            run.add_text('☐' if not is_checked else '☑')
            run.font.name = 'Segoe UI Symbol'
            run.font.size = Pt(12)
            
            # Add space and remaining text
            if checkbox_text:
                paragraph.add_run(' ')
                self._add_formatted_text_single(paragraph, checkbox_text)
            
            # Apply list paragraph style
            self.style_mapper.apply_style_safely(paragraph, self.style_mapper.get_style("list_paragraph"))
            return
        
        # Split text by formatting markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                # Italic text
                italic_text = part[1:-1]
                run = paragraph.add_run(italic_text)
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                # Code text
                code_text = part[1:-1]
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
            else:
                # Regular text - clean up any remaining formatting
                remaining_text = part
                remaining_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', remaining_text)
                remaining_text = re.sub(r'\*([^*]+)\*', r'\1', remaining_text)
                if remaining_text:
                    paragraph.add_run(remaining_text)
