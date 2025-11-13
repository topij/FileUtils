"""DOCX Template Manager for FileUtils.

Manages DOCX templates, style mappings, and template configuration.
"""

import os
from pathlib import Path
from typing import Any, Dict, Optional, Union

import yaml

from ..utils.logging import setup_logger
from ..utils.pathing import find_project_root


class DocxTemplateManager:
    """Manages DOCX templates and their configurations."""

    def __init__(self, config: Dict[str, Any], project_root: Optional[Path] = None):
        """Initialize template manager.

        Args:
            config: Configuration dictionary containing template settings
            project_root: Project root directory for template resolution
        """
        self.config = config
        self.project_root = project_root or self._find_project_root()
        self.logger = setup_logger(self.__class__.__name__)
        self.template_config = self._load_template_config()
        self.style_mappings = self._load_style_mappings()

    def _load_template_config(self) -> Dict[str, Any]:
        """Load template configuration from config or defaults."""
        template_config = self.config.get("docx_templates", {})

        # Default template configuration
        default_config = {
            "template_dir": "templates",
            "default_template": "style-template-doc.docx",
            "templates": {
                "default": "style-template-doc.docx",  # Generic template for sharing
                "review": "style-template-doc.docx",
                "report": "style-template-doc.docx",
                "ip_template": "IP-template-doc.docx",  # Personal IP template
            },
            "fallback_template": None,  # Use python-docx default
        }

        # Merge with provided config
        for key, value in template_config.items():
            if key in default_config:
                if isinstance(default_config[key], dict) and isinstance(value, dict):
                    default_config[key].update(value)
                else:
                    default_config[key] = value
            else:
                default_config[key] = value

        return default_config

    def _load_style_mappings(self) -> Dict[str, str]:
        """Load style mappings from config or defaults."""
        style_config = self.config.get("style_mapping", {})

        # Default style mappings
        default_mappings = {
            "table": "IP-table_light",
            "table_fallback": "IP-table",
            "table_default": "Table Grid",
            "heading_1": "Heading 1",
            "heading_2": "Heading 2",
            "heading_3": "Heading 3",
            "list_bullet": "List Bullet",
            "list_number": "List Number",
            "list_paragraph": "List Paragraph",
            "normal": "Normal",
        }

        # Merge with provided mappings
        default_mappings.update(style_config)
        return default_mappings

    def _find_project_root(self) -> Optional[Path]:
        """Find project root directory using shared helper."""
        return find_project_root()

    def get_template_path(self, template_name: Optional[str] = None) -> Optional[Path]:
        """Get the path to a specific template.

        Args:
            template_name: Name of the template to use

        Returns:
            Path to template file or None if not found
        """
        if not template_name:
            template_name = self.template_config.get("default_template")

        # Get template filename
        template_filename = self.template_config.get("templates", {}).get(template_name)

        # If not found in configuration, check if it's a filename directly
        if not template_filename:
            # Check if template_name is actually a filename
            if template_name.endswith(".docx"):
                template_filename = template_name
            else:
                # Try with .docx extension
                template_filename = f"{template_name}.docx"

            # Verify the file exists before proceeding
            template_dir = self.template_config.get("template_dir", "templates")
            test_paths = [
                Path(template_dir) / template_filename,
                Path("templates") / template_filename,
                Path("data/templates") / template_filename,
                Path(template_filename),
            ]

            if not any(path.exists() for path in test_paths):
                self.logger.warning(
                    f"Template '{template_name}' not found in configuration"
                )
                return None

        # Look for template in multiple locations
        template_dir = self.template_config.get("template_dir", "templates")

        # Build search paths, using project root if available
        search_paths = []

        if self.project_root:
            # Use project root for absolute paths
            search_paths.extend(
                [
                    self.project_root / template_dir / template_filename,
                    self.project_root / "templates" / template_filename,
                    self.project_root / "data/templates" / template_filename,
                    self.project_root / "src/conversion/templates" / template_filename,
                ]
            )

        # Add relative paths as fallback
        search_paths.extend(
            [
                Path(template_dir) / template_filename,
                Path("templates") / template_filename,
                Path("data/templates") / template_filename,
                Path(template_filename),  # Current directory
                Path("src/conversion/templates")
                / template_filename,  # Your old location
            ]
        )

        for template_path in search_paths:
            if template_path.exists():
                self.logger.debug(f"Found template: {template_path}")
                return template_path

        self.logger.warning(
            f"Template file '{template_filename}' not found in any location"
        )
        return None

    def get_style_name(self, style_type: str) -> str:
        """Get the style name for a specific style type.

        Args:
            style_type: Type of style (e.g., 'table', 'heading_1')

        Returns:
            Style name to use in DOCX
        """
        return self.style_mappings.get(style_type, style_type)

    def get_table_style_with_fallback(self) -> str:
        """Get table style with fallback chain.

        Returns:
            First available table style
        """
        fallback_chain = [
            self.style_mappings.get("table"),
            self.style_mappings.get("table_fallback"),
            self.style_mappings.get("table_default"),
            "Table Grid",  # Final fallback
        ]

        for style in fallback_chain:
            if style:
                return style

        return "Table Grid"  # Ultimate fallback

    def validate_template(self, template_path: Path) -> bool:
        """Validate that a template file exists and is readable.

        Args:
            template_path: Path to template file

        Returns:
            True if template is valid, False otherwise
        """
        try:
            if not template_path.exists():
                return False

            # Try to open with python-docx to validate
            from docx import Document

            doc = Document(template_path)
            return True
        except Exception as e:
            self.logger.error(f"Template validation failed for {template_path}: {e}")
            return False

    def list_available_templates(self) -> Dict[str, Path]:
        """List all available templates.

        Returns:
            Dictionary mapping template names to their paths
        """
        available = {}
        template_dir = self.template_config.get("template_dir", "templates")

        # Check configured templates
        for name, filename in self.template_config.get("templates", {}).items():
            template_path = self.get_template_path(name)
            if template_path:
                available[name] = template_path

        # Also scan template directory for additional templates (as fallback)
        template_dir_path = Path(template_dir)
        if template_dir_path.exists():
            for template_file in template_dir_path.glob("*.docx"):
                if template_file not in available.values():
                    # Use filename as template name for discovered templates
                    name = template_file.stem
                    available[name] = template_file

        return available

    def get_template_info(self, template_name: Optional[str] = None) -> Dict[str, Any]:
        """Get information about a template.

        Args:
            template_name: Name of template to get info for

        Returns:
            Dictionary with template information
        """
        template_path = self.get_template_path(template_name)
        if not template_path:
            return {"error": f"Template '{template_name}' not found"}

        try:
            from docx import Document

            doc = Document(template_path)

            # Get available styles
            available_styles = [style.name for style in doc.styles]

            # Check for headers and footers
            header_footer_info = self._get_header_footer_info(doc)

            return {
                "name": template_name or "default",
                "path": str(template_path),
                "exists": True,
                "available_styles": available_styles,
                "style_count": len(available_styles),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "headers_footers": header_footer_info,
            }
        except Exception as e:
            return {
                "name": template_name or "default",
                "path": str(template_path),
                "exists": True,
                "error": str(e),
            }

    def _get_header_footer_info(self, doc) -> Dict[str, Any]:
        """Get information about headers and footers in the document.

        Args:
            doc: Document object

        Returns:
            Dictionary with header/footer information
        """
        try:
            info = {
                "has_headers": False,
                "has_footers": False,
                "header_count": 0,
                "footer_count": 0,
                "header_types": [],
                "footer_types": [],
            }

            # Check for headers
            for section in doc.sections:
                if section.header:
                    info["has_headers"] = True
                    info["header_count"] += 1
                    if section.header.is_linked_to_previous:
                        info["header_types"].append("linked")
                    else:
                        info["header_types"].append("unique")

                if section.footer:
                    info["has_footers"] = True
                    info["footer_count"] += 1
                    if section.footer.is_linked_to_previous:
                        info["footer_types"].append("linked")
                    else:
                        info["footer_types"].append("unique")

            return info
        except Exception as e:
            return {"error": f"Failed to analyze headers/footers: {e}"}
