"""Local filesystem storage implementation."""

import json
import warnings
from pathlib import Path
from typing import Any, Dict, Optional, Union

import pandas as pd
import yaml

from ..core.base import BaseStorage, StorageOperationError
from ..utils.common import ensure_path
from ..utils.dataframe_io import (
    dataframe_to_json,
    dataframe_to_yaml,
    json_to_dataframe,
    read_csv_with_inference,
    yaml_to_dataframe,
)


class LocalStorage(BaseStorage):
    """Local filesystem storage implementation."""

    def save_dataframe(
        self, df: pd.DataFrame, file_path: Union[str, Path], **kwargs
    ) -> str:
        """Save DataFrame to local filesystem.

        Args:
            df: DataFrame to save
            file_path: Path to save to
            **kwargs: Additional arguments for saving:
                - sheet_name: Sheet name for Excel files
                - orient: Orientation for JSON files ("records", "index", etc.)
                - yaml_options: Dict of options for yaml.safe_dump
                - compression: Compression options for parquet files

        Returns:
            String path where the file was saved
        """
        try:
            path = ensure_path(file_path)
            suffix = path.suffix.lower()

            # Ensure parent directory exists
            path.parent.mkdir(parents=True, exist_ok=True)

            if suffix == ".csv":
                df.to_csv(
                    path,
                    index=False,
                    encoding=self.config["encoding"],
                    sep=self.config["csv_delimiter"],
                )
            elif suffix == ".parquet":
                compression = kwargs.get("compression", "snappy")
                df.to_parquet(path, index=False, compression=compression)
            elif suffix in (".xlsx", ".xls"):
                sheet_name = kwargs.get("sheet_name", "Sheet1")
                with pd.ExcelWriter(path, engine="openpyxl") as writer:
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            elif suffix == ".json":
                orient = kwargs.get("orient", "records")
                dataframe_to_json(path, df, orient=orient, indent=2)
            elif suffix == ".yaml" or suffix == ".yml":
                yaml_options = kwargs.get("yaml_options", {})
                orient = kwargs.get("orient", "records")
                dataframe_to_yaml(
                    path,
                    df,
                    orient=orient,
                    yaml_options=yaml_options,
                    encoding=self.config["encoding"],
                )
            else:
                raise ValueError(f"Unsupported file format: {suffix}")

            return str(path)
        except yaml.YAMLError as e:
            raise StorageOperationError(
                f"Failed to save YAML file - YAML error: {e}"
            ) from e
        except Exception as e:
            raise StorageOperationError(f"Failed to save DataFrame: {e}") from e

    def load_dataframe(self, file_path: Union[str, Path], **kwargs) -> pd.DataFrame:
        """Load DataFrame from local filesystem."""
        try:
            path = ensure_path(file_path)
            suffix = path.suffix.lower()

            if suffix == ".csv":
                return self._load_csv_with_inference(path)
            elif suffix == ".parquet":
                return pd.read_parquet(path)
            elif suffix in (".xlsx", ".xls"):
                return pd.read_excel(path, engine="openpyxl")
            elif suffix == ".json":
                return self._load_json_as_dataframe(path)
            elif suffix == ".yaml":
                return self._load_yaml_as_dataframe(path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")

        except Exception as e:
            raise StorageOperationError(f"Failed to load DataFrame: {e}") from e

    def _load_csv_with_inference(self, path: Path) -> pd.DataFrame:
        """Load CSV with delimiter inference."""
        return read_csv_with_inference(
            path=path,
            encoding=self.config["encoding"],
            quoting=self.config["quoting"],
            fallback_sep=self.config["csv_delimiter"],
        )

    def _load_json_as_dataframe(self, path: Path) -> pd.DataFrame:
        """Load JSON file as DataFrame.

        Supports both list of records and dictionary formats.
        """
        try:
            return json_to_dataframe(path, self.config["encoding"])
        except Exception as e:
            raise StorageOperationError(f"Failed to load JSON file: {e}") from e

    def _load_yaml_as_dataframe(self, path: Path) -> pd.DataFrame:
        """Load YAML file as DataFrame.

        Supports both list of records and dictionary formats.
        Handles YAML-specific errors separately for better error messages.
        """
        try:
            return yaml_to_dataframe(path, self.config["encoding"])
        except Exception as e:
            raise StorageOperationError(f"Failed to load YAML file: {e}") from e

    def exists(self, file_path: Union[str, Path]) -> bool:
        """Check if file exists."""
        return Path(file_path).exists()

    def delete(self, file_path: Union[str, Path]) -> bool:
        """Delete file from local filesystem."""
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                return True
            return False
        except Exception as e:
            raise StorageOperationError(
                f"Failed to delete file from local filesystem: {e}"
            ) from e

    def save_bytes(self, content: bytes, file_path: Union[str, Path]) -> str:
        """Save raw bytes to a local file path."""
        try:
            path = ensure_path(file_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            with open(path, "wb") as f:
                f.write(content)
            return str(path)
        except Exception as e:
            raise StorageOperationError(f"Failed to save bytes: {e}") from e

    def load_yaml(self, file_path: Union[str, Path], **kwargs) -> Any:
        """Load YAML file from local filesystem."""
        try:
            path = Path(file_path)

            # If the exact file doesn't exist, try to find a file with timestamp
            if not path.exists():
                # Look for files matching the pattern (with timestamp)
                pattern = f"{path.stem}_*{path.suffix}"
                matching_files = list(path.parent.glob(pattern))

                if matching_files:
                    # Use the most recent file (by modification time)
                    path = max(matching_files, key=lambda f: f.stat().st_mtime)
                else:
                    # If no timestamped file found, raise the original error
                    raise FileNotFoundError(f"File not found: {path}")

            if path.suffix.lower() not in (".yaml", ".yml"):
                raise ValueError("File must have .yaml or .yml extension")

            with open(path, "r", encoding=self.config["encoding"]) as f:
                return yaml.safe_load(f)
        except Exception as e:
            raise StorageOperationError(f"Failed to load YAML file: {e}") from e

    def load_json(self, file_path: Union[str, Path], **kwargs) -> Any:
        """Load JSON file from local filesystem."""
        try:
            path = Path(file_path)

            # If the exact file doesn't exist, try to find a file with timestamp
            if not path.exists():
                # Look for files matching the pattern (with timestamp)
                pattern = f"{path.stem}_*{path.suffix}"
                matching_files = list(path.parent.glob(pattern))

                if matching_files:
                    # Use the most recent file (by modification time)
                    path = max(matching_files, key=lambda f: f.stat().st_mtime)
                else:
                    # If no timestamped file found, raise the original error
                    raise FileNotFoundError(f"File not found: {path}")

            if path.suffix.lower() != ".json":
                raise ValueError("File must have .json extension")

            with open(path, "r", encoding=self.config["encoding"]) as f:
                return json.load(f, **kwargs)
        except Exception as e:
            raise StorageOperationError(f"Failed to load JSON file: {e}") from e

    def save_dataframes(
        self,
        dataframes: Dict[str, pd.DataFrame],
        file_path: Union[str, Path],
        file_format: Optional[str] = None,
        **kwargs,
    ) -> Dict[str, str]:
        """Save multiple DataFrames to local filesystem.

        Args:
            dataframes: Dictionary of DataFrames to save
            file_path: Path to save to
            file_format: File format to save as
            **kwargs: Additional arguments for saving (e.g., engine for Excel)

        Returns:
            Dictionary mapping sheet names to saved file paths. For Excel files,
            all sheets will map to the same file path.
        """
        try:
            path = ensure_path(file_path)

            # Ensure parent directory exists
            path.parent.mkdir(parents=True, exist_ok=True)

            inferred_ext = path.suffix.lstrip(".").lower()
            fmt = inferred_ext

            if file_format is not None and file_format.lower() != inferred_ext:
                warnings.warn(
                    "save_dataframes(file_format=...) is deprecated; format is inferred from file_path",
                    DeprecationWarning,
                    stacklevel=2,
                )

            if fmt in ("xlsx", "xls"):
                # Save all DataFrames to a single Excel file
                with pd.ExcelWriter(
                    path, engine=kwargs.get("engine", "openpyxl")
                ) as writer:
                    for sheet_name, df in dataframes.items():
                        # Handle MultiIndex columns by flattening them
                        if isinstance(df.columns, pd.MultiIndex):
                            df_to_save = df.copy()
                            df_to_save.columns = [
                                "_".join(col).strip() if isinstance(col, tuple) else col
                                for col in df_to_save.columns
                            ]
                        else:
                            df_to_save = df

                        # Use index=True for MultiIndex columns, otherwise use index=False
                        include_index = isinstance(
                            df.columns, pd.MultiIndex
                        ) or isinstance(df.index, pd.MultiIndex)
                        df_to_save.to_excel(
                            writer, sheet_name=sheet_name, index=include_index
                        )
                # For Excel files, return mapping of sheet names to file path
                return {sheet_name: str(path) for sheet_name in dataframes.keys()}
            else:
                # Save each DataFrame to a separate file
                saved_files = {}
                for sheet_name, df in dataframes.items():
                    # Create unique file name for each sheet
                    sheet_path = path.parent / f"{path.stem}_{sheet_name}.{fmt}"
                    saved_path = self.save_dataframe(df, sheet_path, **kwargs)
                    saved_files[sheet_name] = saved_path
                return saved_files

        except Exception as e:
            raise StorageOperationError(f"Failed to save DataFrames: {e}") from e

    def save_document(
        self,
        content: Union[str, Dict[str, Any], bytes, Path],
        file_path: Union[str, Path],
        file_type: str,
        **kwargs,
    ) -> str:
        """Save document content to local filesystem.

        Args:
            content: Document content (string, dict, bytes, or Path).
                    For PPTX: accepts bytes (file content) or Path/str (path to source .pptx file).
            file_path: Path to save to
            file_type: Type of document (docx, md, pdf, pptx)
            **kwargs: Additional arguments for saving

        Returns:
            String path where the file was saved
        """
        try:
            path = ensure_path(file_path)
            suffix = path.suffix.lower()

            # Ensure parent directory exists
            path.parent.mkdir(parents=True, exist_ok=True)

            if suffix == ".docx":
                if not isinstance(content, (str, dict)):
                    raise ValueError("DOCX content must be string or dict")
                return self._save_docx(content, path, **kwargs)
            elif suffix == ".md":
                if not isinstance(content, (str, dict)):
                    raise ValueError("Markdown content must be string or dict")
                return self._save_markdown(content, path, **kwargs)
            elif suffix == ".pdf":
                if not isinstance(content, (str, dict)):
                    raise ValueError("PDF content must be string or dict")
                return self._save_pdf(content, path, **kwargs)
            elif suffix == ".pptx":
                if not isinstance(content, (bytes, Path, str)):
                    raise ValueError("PPTX content must be bytes, Path, or str")
                return self._save_pptx(content, path, **kwargs)
            elif suffix == ".json":
                if not isinstance(content, (str, dict)):
                    raise ValueError("JSON content must be string or dict")
                return self._save_json(content, path, **kwargs)
            elif suffix in (".yaml", ".yml"):
                if not isinstance(content, (str, dict)):
                    raise ValueError("YAML content must be string or dict")
                return self._save_yaml(content, path, **kwargs)
            else:
                raise ValueError(f"Unsupported document format: {suffix}")

        except Exception as e:
            raise StorageOperationError(f"Failed to save document: {e}") from e

    def load_document(
        self, file_path: Union[str, Path], **kwargs
    ) -> Union[str, Dict[str, Any], bytes]:
        """Load document content from local filesystem.

        Args:
            file_path: Path to file
            **kwargs: Additional arguments for loading

        Returns:
            Document content (string, dict, or bytes depending on file type).
            For PPTX: returns bytes.
        """
        try:
            path = ensure_path(file_path)
            suffix = path.suffix.lower()

            if suffix == ".docx":
                return self._load_docx(path, **kwargs)
            elif suffix == ".md":
                return self._load_markdown(path, **kwargs)
            elif suffix == ".pdf":
                return self._load_pdf(path, **kwargs)
            elif suffix == ".pptx":
                return self._load_pptx(path, **kwargs)
            elif suffix == ".json":
                return self._load_json(path, **kwargs)
            elif suffix in (".yaml", ".yml"):
                return self._load_yaml(path, **kwargs)
            else:
                raise ValueError(f"Unsupported document format: {suffix}")

        except Exception as e:
            raise StorageOperationError(f"Failed to load document: {e}") from e

    def _save_docx(
        self, content: Union[str, Dict[str, Any]], path: Path, **kwargs
    ) -> str:
        """Save content to DOCX format using python-docx with template support."""
        try:
            from docx import Document
        except ImportError:
            raise StorageOperationError(
                "python-docx not installed. Install with: pip install 'FileUtils[documents]'"
            )

        # Check if this is markdown content that should be converted
        if isinstance(content, str) and self._is_markdown_content(content):
            return self._save_markdown_as_docx(content, path, **kwargs)

        # Check if template is specified, otherwise use default template
        template_name = kwargs.get("template")
        if template_name is None:
            # Use default template if no template specified
            try:
                from ..templates import DocxTemplateManager

                template_manager = DocxTemplateManager(self.config)
                default_template = template_manager.template_config.get(
                    "default_template"
                )
                if default_template:
                    template_name = "default"  # Use the default template name
            except ImportError:
                pass  # Fall back to no template if template system not available

        if template_name:
            return self._save_with_template(content, path, template_name, **kwargs)

        # Fallback: Default DOCX creation without template
        doc = Document()

        if isinstance(content, str):
            # Simple text content
            doc.add_paragraph(content)
        elif isinstance(content, dict):
            # Structured content
            if "title" in content:
                doc.add_heading(content["title"], 0)

            if "sections" in content:
                for section in content["sections"]:
                    if "heading" in section:
                        doc.add_heading(
                            section["heading"], level=section.get("level", 1)
                        )
                    if "text" in section:
                        doc.add_paragraph(section["text"])
                    if "table" in section and isinstance(section["table"], list):
                        # Add table
                        table = doc.add_table(
                            rows=len(section["table"]), cols=len(section["table"][0])
                        )
                        for i, row in enumerate(section["table"]):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = str(cell)
        else:
            # Fallback: convert to string
            doc.add_paragraph(str(content))

        doc.save(str(path))
        return str(path)

    def _is_markdown_content(self, content: str) -> bool:
        """Check if content looks like markdown."""
        # Simple heuristics to detect markdown
        markdown_indicators = [
            content.startswith("#"),
            "##" in content,
            "###" in content,
            "- " in content,
            "* " in content,
            (
                "|" in content and "|" in content.split("\n")[0]
                if "\n" in content
                else False
            ),
            "**" in content,
            "`" in content,
        ]
        return any(markdown_indicators)

    def _save_markdown_as_docx(
        self, markdown_content: str, path: Path, **kwargs
    ) -> str:
        """Save markdown content as DOCX using template system."""
        try:
            from ..templates import (
                DocxTemplateManager,
                MarkdownToDocxConverter,
                StyleMapper,
            )

            # Initialize template system
            template_manager = DocxTemplateManager(self.config)
            style_mapper = StyleMapper()
            converter = MarkdownToDocxConverter(template_manager, style_mapper)

            # Get conversion options
            template_name = kwargs.get("template")
            add_provenance = kwargs.get("add_provenance", True)
            add_reviewer_instructions = kwargs.get("add_reviewer_instructions", False)
            source_file = kwargs.get("source_file")

            # Convert markdown to DOCX
            doc = converter.convert_markdown_to_docx(
                markdown_content=markdown_content,
                template_name=template_name,
                add_provenance=add_provenance,
                add_reviewer_instructions=add_reviewer_instructions,
                source_file=source_file,
            )

            # Save document
            doc.save(str(path))  # type: ignore
            return str(path)

        except ImportError as e:
            # Fallback to simple conversion if template system not available
            self.logger.warning(
                f"Template system not available, using simple conversion: {e}"
            )
            try:
                from docx import Document
            except ImportError:
                raise StorageOperationError(
                    "python-docx not installed. Install with: pip install 'FileUtils[documents]'"
                ) from e
            doc = Document()
            doc.add_paragraph(markdown_content)
            doc.save(str(path))
            return str(path)
        except Exception as e:
            raise StorageOperationError(
                f"Failed to convert markdown to DOCX: {e}"
            ) from e

    def _save_with_template(
        self,
        content: Union[str, Dict[str, Any]],
        path: Path,
        template_name: str,
        **kwargs,
    ) -> str:
        """Save content using a specific template."""
        try:
            from docx import Document

            from ..templates import DocxTemplateManager

            template_manager = DocxTemplateManager(self.config)
            template_path = template_manager.get_template_path(template_name)

            if template_path and template_path.exists():
                # Load template
                doc = Document(str(template_path))

                # Clear template content
                self._clear_template_content(doc)

                # Add content
                if isinstance(content, str):
                    doc.add_paragraph(content)
                elif isinstance(content, dict):
                    if "title" in content:
                        doc.add_heading(content["title"], 0)
                    if "sections" in content:
                        for section in content["sections"]:
                            if "heading" in section:
                                doc.add_heading(
                                    section["heading"], level=section.get("level", 1)
                                )
                            if "text" in section:
                                doc.add_paragraph(section["text"])
                            if "table" in section and isinstance(
                                section["table"], list
                            ):
                                table = doc.add_table(
                                    rows=len(section["table"]),
                                    cols=len(section["table"][0]),
                                )
                                for i, row in enumerate(section["table"]):
                                    for j, cell in enumerate(row):
                                        table.cell(i, j).text = str(cell)

                doc.save(str(path))
                return str(path)
            else:
                # Fallback to default if template not found
                self.logger.warning(
                    f"Template '{template_name}' not found, using default"
                )
                return self._save_docx(
                    content,
                    path,
                    **{k: v for k, v in kwargs.items() if k != "template"},
                )

        except Exception as e:
            raise StorageOperationError(
                f"Failed to save with template '{template_name}': {e}"
            ) from e

    def _clear_template_content(self, doc):
        """Clear template content while preserving styles, headers, and footers."""
        # Remove all paragraphs from main document body only
        while len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

        # Remove all tables from main document body only
        while len(doc.tables) > 0:
            t = doc.tables[0]._element
            t.getparent().remove(t)

        # Clear document body while preserving headers/footers
        try:
            body = doc._body
            for element in list(body):
                if element.tag.endswith("p") or element.tag.endswith("tbl"):
                    body.remove(element)
        except Exception:
            pass

        # Note: Headers and footers are preserved automatically by python-docx
        # They are stored separately from the main document body

    def _load_docx(self, path: Path, **kwargs) -> str:
        """Load DOCX file and extract text content."""
        try:
            from docx import Document
        except ImportError:
            raise StorageOperationError(
                "python-docx not installed. Install with: pip install 'FileUtils[documents]'"
            )

        doc = Document(str(path))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        return "\n".join(text_content)

    def _save_markdown(
        self, content: Union[str, Dict[str, Any]], path: Path, **kwargs
    ) -> str:
        """Save content to Markdown format."""
        if isinstance(content, str):
            markdown_content = content
        elif isinstance(content, dict):
            # Handle structured content with YAML frontmatter
            frontmatter = content.get("frontmatter", {})
            body = content.get("body", "")

            if frontmatter:
                import yaml

                frontmatter_yaml = yaml.safe_dump(frontmatter, default_flow_style=False)
                markdown_content = f"---\n{frontmatter_yaml}---\n\n{body}"
            else:
                markdown_content = body
        else:
            markdown_content = str(content)

        with open(path, "w", encoding=self.config["encoding"]) as f:
            f.write(markdown_content)

        return str(path)

    def _load_markdown(self, path: Path, **kwargs) -> Union[str, Dict[str, Any]]:
        """Load Markdown file with optional YAML frontmatter."""
        with open(path, "r", encoding=self.config["encoding"]) as f:
            content = f.read()

        # Check for YAML frontmatter
        if content.startswith("---\n"):
            try:
                import yaml

                parts = content.split("---\n", 2)
                if len(parts) >= 3:
                    frontmatter = yaml.safe_load(parts[1])
                    body = parts[2].strip()
                    return {"frontmatter": frontmatter or {}, "body": body}
            except Exception:
                # If frontmatter parsing fails, return as plain text
                pass

        return content

    def _save_pdf(
        self, content: Union[str, Dict[str, Any]], path: Path, **kwargs
    ) -> str:
        """Save content to PDF format using PyMuPDF."""
        try:
            import fitz  # type: ignore # PyMuPDF
        except ImportError:
            raise StorageOperationError(
                "PyMuPDF not installed. Install with: pip install 'FileUtils[documents]'"
            )

        doc = fitz.open()  # Create new PDF
        page = doc.new_page()

        if isinstance(content, str):
            # Simple text content
            page.insert_text((50, 50), content, fontsize=12)
        elif isinstance(content, dict):
            # Structured content
            y_position = 50
            if "title" in content:
                page.insert_text((50, y_position), content["title"], fontsize=16)
                y_position += 30

            if "sections" in content:
                for section in content["sections"]:
                    if "heading" in section:
                        page.insert_text(
                            (50, y_position), section["heading"], fontsize=14
                        )
                        y_position += 25
                    if "text" in section:
                        page.insert_text((50, y_position), section["text"], fontsize=12)
                        y_position += 20
        else:
            # Fallback: convert to string
            page.insert_text((50, 50), str(content), fontsize=12)

        doc.save(path)
        doc.close()
        return str(path)

    def _save_pptx(self, content: Union[bytes, Path, str], path: Path, **kwargs) -> str:
        """Save PPTX content to local filesystem.

        Args:
            content: PPTX file content as bytes, or Path/str to source .pptx file
            path: Destination path where file will be saved
            **kwargs: Additional arguments (not used for PPTX)

        Returns:
            String path where the file was saved

        Raises:
            StorageOperationError: If content type is invalid or file operations fail
        """
        try:
            if isinstance(content, bytes):
                # Write bytes directly to file
                with open(path, "wb") as f:
                    f.write(content)
            elif isinstance(content, Path):
                # Copy file from source path (Path object)
                if not content.exists():
                    raise StorageOperationError(
                        f"Source PPTX file not found: {content}"
                    )
                # Read and write in binary mode
                with open(content, "rb") as src:
                    with open(path, "wb") as dst:
                        dst.write(src.read())
            elif isinstance(content, str):
                # String could be a file path - check if it's a valid path to an existing file
                source_path = ensure_path(content)
                if not source_path.exists():
                    # Not a valid file path, raise error about invalid content type
                    raise StorageOperationError(
                        f"Invalid content type for PPTX: str (not a valid file path). "
                        f"Expected bytes or Path/str to an existing source .pptx file. "
                        f"Received string that doesn't point to an existing file: {content}"
                    )
                # Valid file path - copy the file
                with open(source_path, "rb") as src:
                    with open(path, "wb") as dst:
                        dst.write(src.read())
            else:
                raise StorageOperationError(
                    f"Invalid content type for PPTX: {type(content)}. "
                    f"Expected bytes or Path/str to source .pptx file."
                )

            return str(path)
        except Exception as e:
            raise StorageOperationError(f"Failed to save PPTX file: {e}") from e

    def _load_pdf(self, path: Path, **kwargs) -> str:
        """Load PDF file and extract text content."""
        try:
            import fitz  # type: ignore # PyMuPDF
        except ImportError:
            raise StorageOperationError(
                "PyMuPDF not installed. Install with: pip install 'FileUtils[documents]'"
            )

        doc = fitz.open(path)
        text_content = []

        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_content.append(text)

        doc.close()
        return "\n\n".join(text_content)

    def _load_pptx(self, path: Path, **kwargs) -> bytes:
        """Load PPTX file as bytes.

        Args:
            path: Path to PPTX file
            **kwargs: Additional arguments (not used for PPTX)

        Returns:
            Bytes content of the PPTX file
        """
        try:
            with open(path, "rb") as f:
                return f.read()
        except Exception as e:
            raise StorageOperationError(f"Failed to load PPTX file: {e}") from e

    def _load_json(self, path: Path, **kwargs) -> Dict[str, Any]:
        """Load JSON file."""
        try:
            import json

            with open(path, "r", encoding=self.config["encoding"]) as f:
                return json.load(f, **kwargs)
        except Exception as e:
            raise StorageOperationError(f"Failed to load JSON file: {e}") from e

    def _load_yaml(self, path: Path, **kwargs) -> Dict[str, Any]:
        """Load YAML file."""
        try:
            import yaml

            with open(path, "r", encoding=self.config["encoding"]) as f:
                return yaml.safe_load(f, **kwargs)
        except Exception as e:
            raise StorageOperationError(f"Failed to load YAML file: {e}") from e

    def _save_json(
        self, content: Union[str, Dict[str, Any]], path: Path, **kwargs
    ) -> str:
        """Save content as JSON file."""
        try:
            import json

            # Custom JSON encoder to handle pandas types
            class PandasJSONEncoder(json.JSONEncoder):
                def default(self, obj):
                    if hasattr(obj, "isoformat"):  # datetime, Timestamp
                        return obj.isoformat()
                    elif hasattr(obj, "tolist"):  # numpy arrays (check this first)
                        return obj.tolist()
                    elif hasattr(obj, "item"):  # numpy scalar types
                        return obj.item()
                    return super().default(obj)

            with open(path, "w", encoding=self.config["encoding"]) as f:
                json.dump(
                    content,
                    f,
                    indent=kwargs.get("indent", 2),
                    cls=PandasJSONEncoder,
                    **kwargs,
                )
            return str(path)
        except Exception as e:
            raise StorageOperationError(f"Failed to save JSON file: {e}") from e

    def _save_yaml(
        self, content: Union[str, Dict[str, Any]], path: Path, **kwargs
    ) -> str:
        """Save content as YAML file."""
        try:
            import yaml

            with open(path, "w", encoding=self.config["encoding"]) as f:
                yaml.dump(content, f, default_flow_style=False, **kwargs)
            return str(path)
        except Exception as e:
            raise StorageOperationError(f"Failed to save YAML file: {e}") from e

    def list_directory(
        self,
        directory_path: Union[str, Path],
        pattern: Optional[str] = None,
        files_only: bool = False,
        directories_only: bool = False,
    ) -> list:
        """List files and directories in a local filesystem path.

        Args:
            directory_path: Path to directory
            pattern: Optional glob pattern to filter results (e.g., "*.yml", "*.pptx")
            files_only: If True, return only files (exclude directories)
            directories_only: If True, return only directories (exclude files)

        Returns:
            List[str]: List of file/directory names (not full paths).
                     Returns empty list if directory doesn't exist or on error.
        """
        try:
            path = Path(directory_path)
            if not path.exists() or not path.is_dir():
                return []

            # Get all items
            items = []
            if pattern:
                # Use glob pattern matching
                import fnmatch

                for item in path.iterdir():
                    if fnmatch.fnmatch(item.name, pattern):
                        items.append(item)
            else:
                items = list(path.iterdir())

            # Filter by type
            result = []
            for item in items:
                # Skip hidden files/directories
                if item.name.startswith("."):
                    continue

                if files_only and not item.is_file():
                    continue
                if directories_only and not item.is_dir():
                    continue

                # Return just the name (not full path)
                result.append(item.name)

            return sorted(result)
        except Exception:
            # Return empty list on any error (per requirements)
            return []
